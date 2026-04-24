#!/usr/bin/env python3
"""
Сборщик ВКР: markdown-разделы -> ВКР_полный.docx

Оформление по ГОСТ 7.32-2017 / Правила ВКР ПИ (11.2020):
- Times New Roman, 13 пт, межстрочный интервал 1.5
- Поля: левое 3 см, правое 1.5 см, верх/низ 2 см
- Выравнивание текста - по ширине, абзацный отступ 1.25 см
- Заголовки глав: 16 пт, жирный, перед 0, после 12
- Заголовки параграфов: 14 пт, жирный, перед 12, после 6
- Нумерация страниц - внизу по центру
- Таблица: название слева, без отступа, "Таблица X - название"
- Рисунок: подпись снизу по центру, жирный курсив 12 пт
- Каждая глава начинается с нового листа
- После номера раздела/подраздела точка не ставится
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

BASE = Path(__file__).parent
OUTPUT_PATH = BASE / "ВКР_полный.docx"
DRAFTS = BASE / "drafts"
DIAGRAMS = BASE / "diagrams"

# Введение
INTRO_FILE = "00_введение.md"

# Порядок файлов для Главы 1
CHAPTER1_FILES = [
    "01_1_краудфандинг.md",
    "01_2_анализ_рынка.md",
    "01_3_блокчейн.md",
    "01_4_web3.md",
    "01_5_стейкхолдеры.md",
    "01_6_концепция.md",
    "01_7_выводы.md",
]

# Порядок файлов для Главы 2 (бизнес-процесс)
CHAPTER2_FILES = [
    "02_1_as_is.md",
    "02_2_требования.md",
    "02_3_to_be.md",
    "02_4_сравнение.md",
    "02_5_верификация.md",
    "02_6_выводы.md",
]

# Порядок файлов для Главы 3 (архитектура и прототип)
CHAPTER3_FILES = [
    "03_1_архитектура.md",
    "03_2_смарт_контракты.md",
    "03_3_данные.md",
    "03_4_сценарии.md",
    "03_5_план_внедрения.md",
    "03_6_прототип.md",
    "03_7_выводы.md",
]

# Заключение
CONCLUSION_FILE = "04_заключение.md"

# Приложения (по алфавиту: А, Б, В, Г, ...)
APPENDIX_FILES = [
    "app_a_экраны.md",
]

# Рисунки: (номер, файл, подпись, ширина_см, landscape)
# landscape=True → рисунок на отдельной альбомной странице (ширина 25 см)
FIGURES = [
    ("1.1", "crowdfunding_actors.png", "Схема взаимодействия участников краудфандинга", 14, False),
    ("1.2", "market_chart.png", "Динамика глобального и российского рынков краудфандинга", 16, False),
    ("1.3", "ishikawa_a1.png", "Диаграмма Исикавы: ключевые проблемы краудфандинговых платформ", 16, False),
    ("1.4", "guggenberger_matrix.png", "Покрытие принципов проектирования Guggenberger существующими решениями", 16, False),
    ("1.5", "mendelow_matrix.png", "Матрица Mendelow (Power-Interest Grid) заинтересованных сторон", 14, False),
    ("1.6", "concept_architecture.png", "Концептуальная архитектура платформы", 15, False),
    ("2.1", "bpmn_renders/bpmn_as_is_l0.png", "BPMN-модель AS-IS, уровень L1 (Descriptive): описательная модель Web2-краудфандинга", 16, False),
    ("2.2", "bpmn_renders/bpmn_as_is_l1.png", "BPMN-модель AS-IS, уровень L2 (Analytical): распределение ответственности и узкие места", 16, False),
    ("2.3", "use_case_g1.png", "Диаграмма вариантов использования целевой платформы", 15, False),
    ("2.4", "bpmn_renders/bpmn_to_be_l0.png", "BPMN-модель TO-BE, уровень L1 (Descriptive): описательная модель платформы токенизированного инвестирования", 16, False),
    ("2.5", "bpmn_renders/bpmn_to_be_l1.png", "BPMN-модель TO-BE, уровень L2 (Analytical): распределение ответственности по четырём пулам", 16, False),
    ("3.1", "concept_architecture.png", "Концептуальная архитектура платформы", 16, False),
    ("3.2", "component_diagram.png", "Компонентная диаграмма платформы", 16, False),
    ("3.3", "class_diagram_contracts.png", "Диаграмма классов смарт-контрактов платформы", 16, False),
    ("3.4", "state_campaign.png", "Диаграмма состояний кампании", 15, False),
    ("3.5", "state_milestone.png", "Диаграмма состояний этапа", 15, False),
    ("3.6", "er_e1.png", "ER-диаграмма модели данных платформы", 16, False),
    ("3.7", "sequence_invest.png", "Сценарий инвестирования (UML Sequence)", 15, False),
    ("3.8", "proto_home.png", "Главная страница прототипа: каталог кампаний", 14, False),
    ("3.9", "proto_project_active.png", "Страница активной кампании с голосованием по этапу", 14, False),
    # Приложение Г: экранные формы прототипа
    ("А.1", "proto_home.png", "Каталог кампаний (Home)", 14, False),
    ("А.2", "proto_register.png", "Форма регистрации пользователя", 11, False),
    ("А.3", "proto_login.png", "Форма входа в систему", 11, False),
    ("А.4", "proto_project_funding.png", "Страница кампании в состоянии Funding", 14, False),
    ("А.5", "proto_project_active.png", "Страница кампании в состоянии Active с голосованием", 14, False),
    ("А.6", "proto_project_completed.png", "Страница завершённой кампании", 14, False),
    ("А.7", "proto_create.png", "Форма создания новой кампании", 14, False),
    ("А.8", "proto_dashboard.png", "Панель управления автора проектов", 14, False),
    ("А.9", "proto_profile.png", "Личный кабинет пользователя", 14, False),
]

# Размеры шрифтов по презентации НН (2025)
BODY_FONT_SIZE = 14
HEADING1_FONT_SIZE = 16   # Заголовки глав
HEADING2_FONT_SIZE = 14   # Заголовки параграфов
TABLE_FONT_SIZE = 12      # Текст в таблицах (11-12)
FIGURE_CAPTION_SIZE = 12  # Подписи рисунков (11-12)
PAGE_NUM_FONT_SIZE = 12


# --- Форматирование ---

def set_run_font(run, font_name="Times New Roman", font_size=BODY_FONT_SIZE,
                 bold=False, italic=False):
    """Настройка шрифта для run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_paragraph_format(para, font_size=BODY_FONT_SIZE,
                         bold=False, italic=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_indent_cm=1.25, line_spacing=1.5,
                         space_after=0, space_before=0):
    """Настройка форматирования параграфа и всех его run."""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.alignment = alignment
    if first_indent_cm is not None:
        pf.first_line_indent = Cm(first_indent_cm)
    else:
        pf.first_line_indent = Cm(0)
    for run in para.runs:
        set_run_font(run, font_size=font_size, bold=bold, italic=italic)


def _set_style_font(style, font_size, bold=False):
    """Настроить шрифт стиля на Times New Roman."""
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    for attr in ['asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'cstheme']:
        rFonts.attrib.pop(qn('w:' + attr), None)
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn('w:' + attr), 'Times New Roman')


def setup_document_defaults(doc):
    """Times New Roman 13pt по умолчанию + стили Heading 1/2 для оглавления."""
    # Normal
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    for attr in ['asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'cstheme']:
        rFonts.attrib.pop(qn('w:' + attr), None)
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn('w:' + attr), 'Times New Roman')

    # Document defaults
    styles_elem = doc.styles.element
    rPrDefault = styles_elem.find('.//' + qn('w:rPrDefault'))
    if rPrDefault is not None:
        rPr_d = rPrDefault.find(qn('w:rPr'))
        if rPr_d is not None:
            rFonts_d = rPr_d.find(qn('w:rFonts'))
            if rFonts_d is not None:
                for attr in ['asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'cstheme']:
                    rFonts_d.attrib.pop(qn('w:' + attr), None)
                for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
                    rFonts_d.set(qn('w:' + attr), 'Times New Roman')
            sz = rPr_d.find(qn('w:sz'))
            if sz is not None:
                sz.set(qn('w:val'), str(BODY_FONT_SIZE * 2))
            szCs = rPr_d.find(qn('w:szCs'))
            if szCs is not None:
                szCs.set(qn('w:val'), str(BODY_FONT_SIZE * 2))

    # Heading 1: 16pt, жирный, по центру, перед 0, после 12, с новой страницы
    h1 = doc.styles['Heading 1']
    _set_style_font(h1, HEADING1_FONT_SIZE, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True

    # Heading 2: 14pt, жирный, по центру, перед 12, после 6
    h2 = doc.styles['Heading 2']
    _set_style_font(h2, HEADING2_FONT_SIZE, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.page_break_before = False
    h2.paragraph_format.keep_with_next = True


def setup_page_format(doc):
    """A4, поля по презентации НН: лево 3, право 1, верх/низ 2."""
    for section in doc.sections:
        _setup_section_format(section)


def _setup_section_format(section):
    """Применить стандартные поля A4 к секции."""
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def add_page_number_to_footer(section):
    """Добавить поле PAGE в футер секции (внизу по центру)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if not footer.paragraphs[0].text else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = para.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    for r in (run, run2, run3):
        set_run_font(r, font_size=PAGE_NUM_FONT_SIZE)


def add_section_break(doc):
    """Добавить разрыв секции (следующая страница) и вернуть новую секцию.

    ВАЖНО: add_section_break вставляет sectPr в pPr последнего параграфа,
    что фактически делает его описанием ТЕКУЩЕЙ (становящейся первой) секции.
    Чтобы страницы не становились landscape в Google Docs, в этот sectPr нужно
    скопировать pgSz и pgMar из документного sectPr (который описывает ПОСЛЕДНЮЮ).
    """
    last_para = doc.paragraphs[-1]
    pPr = last_para._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)

    # Копируем pgSz, pgMar, cols, docGrid из body/sectPr (чтобы orientation сохранялась).
    # НЕ копируем headerReference/footerReference - их выставим явно в build().
    body_sectPr = doc.element.body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        import copy
        for tag in ('w:pgSz', 'w:pgMar', 'w:cols', 'w:docGrid', 'w:formProt', 'w:textDirection'):
            el = body_sectPr.find(qn(tag))
            if el is not None:
                sectPr.append(copy.deepcopy(el))

    pPr.append(sectPr)
    return doc.sections[-1]


def add_page_break(doc):
    """Разрыв страницы."""
    last = doc.paragraphs[-1]
    run = last.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def add_title_line(doc, text, bold=False, space_before=0, space_after=0):
    """Строка титульного листа: по центру, 13pt, межстрочный 1.5."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_size=BODY_FONT_SIZE, bold=bold)
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return para


def build_title_page(doc):
    """Титульный лист по Приложению 3 (Приложение 8 к Положению НИУ ВШЭ)."""
    add_title_line(doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ")
    add_title_line(doc, "ВЫСШЕГО ОБРАЗОВАНИЯ")
    add_title_line(doc, "\u00abНАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ")
    add_title_line(doc, "\u00abВЫСШАЯ ШКОЛА ЭКОНОМИКИ\u00bb")

    add_title_line(doc, "")
    add_title_line(doc, "Факультет информатики, математики и компьютерных наук")
    add_title_line(doc, "")
    add_title_line(doc, "")

    add_title_line(doc, "Досталев Артем Алексеевич")

    add_title_line(doc, "")

    add_title_line(doc, "РАЗРАБОТКА АЛГОРИТМА И ПРОТОТИПА ТОКЕНИЗИРОВАННОГО ФОНДА", bold=True)
    add_title_line(doc, "КОЛЛЕКТИВНОГО ИНВЕСТИРОВАНИЯ", bold=True)

    add_title_line(doc, "")

    add_title_line(doc, "Выпускная квалификационная работа \u2013 БАКАЛАВРСКАЯ РАБОТА")
    add_title_line(doc, "по направлению подготовки 38.03.05 Бизнес-информатика")
    add_title_line(doc, "образовательная программа \u00abБизнес-информатика\u00bb")

    add_title_line(doc, "")
    add_title_line(doc, "")

    # Научный руководитель - справа
    para = doc.add_paragraph()
    run = para.add_run("Научный руководитель")
    set_run_font(run, font_size=BODY_FONT_SIZE)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_after = Pt(0)

    para2 = doc.add_paragraph()
    run2 = para2.add_run("к.э.н., доцент БКИФР")
    set_run_font(run2, font_size=BODY_FONT_SIZE)
    para2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para2.paragraph_format.line_spacing = 1.5
    para2.paragraph_format.first_line_indent = Cm(0)
    para2.paragraph_format.space_after = Pt(0)

    para3 = doc.add_paragraph()
    run3 = para3.add_run("Солдатова А. О.")
    set_run_font(run3, font_size=BODY_FONT_SIZE)
    para3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para3.paragraph_format.line_spacing = 1.5
    para3.paragraph_format.first_line_indent = Cm(0)
    para3.paragraph_format.space_after = Pt(0)

    add_title_line(doc, "")
    add_title_line(doc, "")

    add_title_line(doc, "Нижний Новгород 2026")

    add_page_break(doc)


# --- Элементы документа ---

def add_chapter_heading(doc, text):
    """Заголовок главы: стиль Heading 1 (16pt, жирный, с новой страницы)."""
    para = doc.add_paragraph(text, style='Heading 1')
    # Принудительно шрифт на runs (стиль может не применяться к run)
    for run in para.runs:
        set_run_font(run, font_size=HEADING1_FONT_SIZE, bold=True)
    return para


def add_section_heading(doc, text):
    """Заголовок параграфа: стиль Heading 2 (14pt, жирный)."""
    para = doc.add_paragraph(text, style='Heading 2')
    for run in para.runs:
        set_run_font(run, font_size=HEADING2_FONT_SIZE, bold=True)
    return para


def add_subsection_heading(doc, text):
    """Заголовок подпараграфа (уровень 3): стиль Heading 3 (14pt, жирный, курсив)."""
    para = doc.add_paragraph(text, style='Heading 3')
    for run in para.runs:
        set_run_font(run, font_size=HEADING2_FONT_SIZE, bold=True, italic=True)
    return para


def add_toc_field(doc):
    """Вставить автоматическое оглавление (TOC field).
    Обновится при открытии в Word (Ctrl+A, F9)."""
    # Заголовок "ОГЛАВЛЕНИЕ"
    para = doc.add_paragraph("ОГЛАВЛЕНИЕ")
    set_paragraph_format(para, font_size=HEADING1_FONT_SIZE, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_indent_cm=None, space_before=0, space_after=12)

    # Поле TOC
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = toc_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    run2 = toc_para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = toc_para.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_separate)

    run4 = toc_para.add_run("[Обновите оглавление: выделите все (Ctrl+A) и нажмите F9]")
    set_run_font(run4, font_size=BODY_FONT_SIZE, italic=True)

    run5 = toc_para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar_end)

    add_page_break(doc)


def add_body_paragraph(doc, text):
    """Абзац основного текста 13pt с поддержкой **жирного**."""
    para = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run)
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(1.25)
    return para


def add_list_item(doc, text):
    """Маркированный пункт списка (тире). Отступ первой строки 1.5 см, слева 2 см."""
    para = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    first = True
    for part in parts:
        prefix = "\u2013 " if first else ""
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(prefix + part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = para.add_run(prefix + part)
            set_run_font(run)
        first = False
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(1.5)
    pf.left_indent = Cm(2)
    return para


def add_table(doc, caption_text, header_line, rows_text):
    """Таблица с подписью: слева, без отступа, "Таблица X - название". Текст 12pt."""
    cap_para = doc.add_paragraph()
    cap_para.add_run(caption_text)
    set_paragraph_format(cap_para, font_size=TABLE_FONT_SIZE,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_indent_cm=None, space_before=6, space_after=3)
    # Не отрывать от следующего
    pPr = cap_para._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))

    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    rows = []
    for row_text in rows_text:
        cells = [c.strip() for c in row_text.split('|')[1:-1]]
        rows.append(cells)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')

    # Повтор заголовка на новой странице
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            set_paragraph_format(para, font_size=TABLE_FONT_SIZE, bold=True,
                                 first_indent_cm=None, line_spacing=1.0,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < num_cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for para in cell.paragraphs:
                    set_paragraph_format(para, font_size=TABLE_FONT_SIZE, bold=False,
                                         first_indent_cm=None, line_spacing=1.0,
                                         alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Отступ после таблицы
    after = doc.add_paragraph()
    after.paragraph_format.line_spacing = 1.5
    after.paragraph_format.space_before = Pt(6)
    after.paragraph_format.space_after = Pt(0)
    return table


def _insert_section_pgSz(last_para, orient, width_mm, height_mm):
    """Вставить sectPr в конец параграфа с заданной ориентацией страницы."""
    from copy import deepcopy
    pPr = last_para._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)

    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(width_mm * 1440 / 25.4)))
    pgSz.set(qn('w:h'), str(int(height_mm * 1440 / 25.4)))
    pgSz.set(qn('w:orient'), orient)
    sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1134')
    pgMar.set(qn('w:right'), '1134')
    pgMar.set(qn('w:bottom'), '1134')
    pgMar.set(qn('w:left'), '1701')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)

    # Копируем cols/docGrid из body/sectPr если есть
    doc = last_para.part.element.getparent() if hasattr(last_para.part.element, 'getparent') else None
    pPr.append(sectPr)


def add_figure(doc, figure_num, filename, caption, width_cm, landscape=False):
    """Рисунок с подписью. При landscape=True оборачивается в отдельную
    альбомную секцию (297×210 мм) — для широких BPMN-диаграмм."""
    img_path = DIAGRAMS / filename
    if not img_path.exists():
        print(f"  [!] Файл {filename} не найден, пропускаю")
        return

    if landscape:
        # Закрыть текущую portrait-секцию перед landscape-страницей
        last = doc.paragraphs[-1]
        _insert_section_pgSz(last, 'portrait', 210, 297)

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.first_line_indent = Cm(0)
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.paragraph_format.space_after = Pt(0)
    run = pic_para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    # Подпись: жирный курсив 12pt
    cap_para = doc.add_paragraph()
    cap_run = cap_para.add_run(f"Рисунок {figure_num} \u2013 {caption}")
    set_run_font(cap_run, font_size=FIGURE_CAPTION_SIZE, bold=True, italic=True)
    pf = cap_para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0

    if landscape:
        # Закрыть landscape-секцию и вернуться к portrait
        _insert_section_pgSz(cap_para, 'landscape', 297, 210)


# --- Парсер markdown ---

def normalize_text(text):
    """Замена markdown-тире на короткое тире."""
    text = text.replace('---', '\u2013')
    text = text.replace('--', '\u2013')
    return text


def parse_markdown(text):
    """Парсит markdown-текст в список элементов."""
    text = normalize_text(text)
    elements = []
    lines = text.strip().split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith('### '):
            title = line[4:].strip()
            # После номера подподраздела точка не ставится: 2.1.1. -> 2.1.1
            title = re.sub(r'^(\d+\.\d+\.\d+)\.\s', r'\1 ', title)
            elements.append(('subheading', title))
            i += 1
            continue

        if line.startswith('## '):
            title = line[3:].strip()
            # После номера подраздела точка не ставится: 1.1. -> 1.1
            title = re.sub(r'^(\d+\.\d+)\.\s', r'\1 ', title)
            elements.append(('heading', title))
            i += 1
            continue

        if line.startswith('Таблица '):
            caption = line
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and '|' in lines[i]:
                header = lines[i]
                i += 2  # пропускаем разделитель ---
                rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip():
                    rows.append(lines[i])
                    i += 1
                elements.append(('table', caption, header, rows))
            continue

        if line.startswith('Рисунок '):
            match = re.match(r'Рисунок ([\wА-Яа-яЁё]+\.\d+)', line)
            if match:
                elements.append(('figure', match.group(1)))
                i += 1
                continue

        if re.match(r'^\d+\.\s', line):
            elements.append(('body', line))
            i += 1
            continue

        if line.startswith('- '):
            elements.append(('list_item', line[2:]))
            i += 1
            continue

        # Многострочный параграф
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() \
                and not lines[i].startswith('#') \
                and not lines[i].startswith('|') \
                and not lines[i].startswith('Таблица ') \
                and not lines[i].startswith('Рисунок ') \
                and not lines[i].startswith('- ') \
                and not re.match(r'^\d+\.\s', lines[i]):
            para_lines.append(lines[i].rstrip())
            i += 1
        elements.append(('body', ' '.join(para_lines)))

    return elements


def render_elements(doc, elements, figure_map):
    """Рендерит список элементов в документ."""
    for elem in elements:
        kind = elem[0]
        if kind == 'heading':
            add_section_heading(doc, elem[1])
        elif kind == 'subheading':
            add_subsection_heading(doc, elem[1])
        elif kind == 'body':
            add_body_paragraph(doc, elem[1])
        elif kind == 'list_item':
            add_list_item(doc, elem[1])
        elif kind == 'table':
            add_table(doc, elem[1], elem[2], elem[3])
        elif kind == 'figure':
            fig_num = elem[1]
            if fig_num in figure_map:
                _, filename, caption, width, *rest = figure_map[fig_num]
                landscape = rest[0] if rest else False
                add_figure(doc, fig_num, filename, caption, width, landscape)
            else:
                para = doc.add_paragraph()
                para.add_run(f"[Рисунок {fig_num}: см. приложение]")
                set_paragraph_format(para, italic=True,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     first_indent_cm=None)


# --- Сборка ---

def build():
    """Основная функция сборки."""
    print("Создаю чистый документ...")
    doc = Document()

    # Убираем пустой параграф по умолчанию
    if doc.paragraphs and not doc.paragraphs[0].text:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    setup_document_defaults(doc)
    setup_page_format(doc)

    # Секция 0: титульник (БЕЗ номера страницы, симметричные поля)
    first_section = doc.sections[0]
    first_section.footer.is_linked_to_previous = False
    first_section.left_margin = Cm(2)
    first_section.right_margin = Cm(2)
    # Пустой футер - номер не показываем

    # Автообновление полей (оглавление) при открытии документа
    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    figure_map = {f[0]: f for f in FIGURES}

    # === ТИТУЛЬНЫЙ ЛИСТ ===
    print("  Генерирую титульный лист...")
    build_title_page(doc)

    # Секция 1: оглавление
    toc_section = add_section_break(doc)
    _setup_section_format(toc_section)

    # === ОГЛАВЛЕНИЕ (автоматическое) ===
    add_toc_field(doc)

    # Секция 2: основной текст
    new_section = add_section_break(doc)
    _setup_section_format(new_section)

    # === ВВЕДЕНИЕ ===
    intro_path = DRAFTS / INTRO_FILE
    if intro_path.exists():
        print(f"  Обрабатываю {INTRO_FILE}...")
        elements = parse_markdown(intro_path.read_text(encoding='utf-8'))
        for elem in elements:
            if elem[0] == 'heading':
                # ВВЕДЕНИЕ - как заголовок главы (16pt), с новой страницы
                add_chapter_heading(doc, elem[1])
            elif elem[0] == 'body':
                add_body_paragraph(doc, elem[1])
            elif elem[0] == 'list_item':
                add_list_item(doc, elem[1])
    else:
        print(f"  [!] {INTRO_FILE} не найден")

    # === ГЛАВА 1 ===
    add_chapter_heading(doc, "ГЛАВА 1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ВЫЯВЛЕНИЕ ПРОБЛЕМ")

    for fname in CHAPTER1_FILES:
        fpath = DRAFTS / fname
        if not fpath.exists():
            print(f"  [!] {fname} не найден")
            continue
        print(f"  Обрабатываю {fname}...")
        elements = parse_markdown(fpath.read_text(encoding='utf-8'))
        render_elements(doc, elements, figure_map)

    # === ГЛАВА 2 ===
    add_chapter_heading(doc, "ГЛАВА 2 ПРОЕКТИРОВАНИЕ ПЛАТФОРМЫ ТОКЕНИЗИРОВАННОГО ИНВЕСТИРОВАНИЯ")

    for fname in CHAPTER2_FILES:
        fpath = DRAFTS / fname
        if not fpath.exists():
            print(f"  [!] {fname} не найден")
            continue
        print(f"  Обрабатываю {fname}...")
        elements = parse_markdown(fpath.read_text(encoding='utf-8'))
        render_elements(doc, elements, figure_map)

    # === ГЛАВА 3 ===
    add_chapter_heading(doc, "ГЛАВА 3 АРХИТЕКТУРА И РЕАЛИЗАЦИЯ ПРОТОТИПА ПЛАТФОРМЫ")

    for fname in CHAPTER3_FILES:
        fpath = DRAFTS / fname
        if not fpath.exists():
            print(f"  [!] {fname} не найден")
            continue
        print(f"  Обрабатываю {fname}...")
        elements = parse_markdown(fpath.read_text(encoding='utf-8'))
        render_elements(doc, elements, figure_map)

    # === ЗАКЛЮЧЕНИЕ ===
    conclusion_path = DRAFTS / CONCLUSION_FILE
    if conclusion_path.exists():
        print(f"  Обрабатываю {CONCLUSION_FILE}...")
        elements = parse_markdown(conclusion_path.read_text(encoding='utf-8'))
        for elem in elements:
            if elem[0] == 'heading':
                add_chapter_heading(doc, elem[1])
            elif elem[0] == 'body':
                add_body_paragraph(doc, elem[1])
            elif elem[0] == 'list_item':
                add_list_item(doc, elem[1])
    else:
        print(f"  [!] {CONCLUSION_FILE} не найден")

    # === ПРИЛОЖЕНИЯ ===
    for app_name in APPENDIX_FILES:
        app_path = DRAFTS / app_name
        if not app_path.exists():
            print(f"  [!] {app_name} не найден")
            continue
        print(f"  Обрабатываю приложение {app_name}...")
        elements = parse_markdown(app_path.read_text(encoding='utf-8'))
        # По ГОСТ/Правилам ВКР ПИ:
        #   1) Новая страница
        #   2) Слово «ПРИЛОЖЕНИЕ <буква>» по центру верха страницы (не жирное, обычный 14 pt)
        #   3) Тематический заголовок — по центру, полужирный, без точки
        heading_count = 0
        for elem in elements:
            if elem[0] == 'heading':
                heading_count += 1
                if heading_count == 1:
                    # «ПРИЛОЖЕНИЕ А» — с новой страницы, по центру, без красной строки
                    add_page_break(doc)
                    p = doc.add_paragraph()
                    run = p.add_run(elem[1])
                    set_run_font(run, font_size=BODY_FONT_SIZE, bold=False)
                    pf = p.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.first_line_indent = Cm(0)
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(6)
                    pf.line_spacing = 1.5
                elif heading_count == 2:
                    # Тематический заголовок: по центру, полужирный, без абзацного отступа, без точки
                    p = doc.add_paragraph()
                    run = p.add_run(elem[1].rstrip('.'))
                    set_run_font(run, font_size=BODY_FONT_SIZE, bold=True)
                    pf = p.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.first_line_indent = Cm(0)
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(12)
                    pf.line_spacing = 1.5
                else:
                    add_section_heading(doc, elem[1])
            elif elem[0] == 'subheading':
                add_subsection_heading(doc, elem[1])
            elif elem[0] == 'body':
                add_body_paragraph(doc, elem[1])
            elif elem[0] == 'list_item':
                add_list_item(doc, elem[1])
            elif elem[0] == 'figure':
                fig_num = elem[1]
                if fig_num in figure_map:
                    _, filename, caption, width, *rest = figure_map[fig_num]
                    landscape = rest[0] if rest else False
                    add_figure(doc, fig_num, filename, caption, width, landscape)

    # Нумерация страниц: применить ко всем секциям КРОМЕ титульника (секция 0)
    # после того как весь контент добавлен, но до save
    for section in list(doc.sections)[1:]:
        add_page_number_to_footer(section)

    # Сохранение
    doc.save(str(OUTPUT_PATH))

    check = Document(str(OUTPUT_PATH))
    print(f"\nГотово: {OUTPUT_PATH.name}")
    print(f"  Параграфов: {len(check.paragraphs)}")
    print(f"  Таблиц: {len(check.tables)}")
    print(f"  Размер: {OUTPUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
